"""
06_newsletter.py
=================
STAGE 6 of the FMCG M&A Newsletter pipeline: NEWSLETTER GENERATION



What this version DOES do:
  1. VISUAL STRUCTURE -- the agreed McKinsey-style design: serif masthead,
     a data-derived Executive Highlights block, a stat row, a Spotlight
     tier (top 3 by disclosed value, full card treatment) followed by a
     flat "Also This Issue" list (the rest, compact), real bordered boxes
     for every item, single gold accent color throughout.
  2. RECENCY WINDOW -- the one new piece of logic in this version: only
     articles published in the last DAYS_LOOKBACK days (15 by default)
     are included. This is a ROLLING window measured from the moment the
     script runs (datetime.now()), the same convention 01_ingest.py uses
     for its own 30-day window, not a fixed calendar month.



INPUT / OUTPUT
--------------
Reads:  data/summarized_news.csv               (written by 05_summarize.py)
        data/raw_news.csv, processed_news.csv,    (read ONLY for the
        relevant_news.csv, credible_news.csv       appendix funnel table)
Writes: data/newsletter_YYYY_MM.docx
"""

import os
import re
from datetime import datetime, timezone, timedelta
from email.utils import parsedate_to_datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

# ─────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────

INPUT_FILE = os.path.join("data", "summarized_news.csv")
OUTPUT_FILE = os.path.join("data", f"newsletter_{datetime.now().strftime('%Y_%m')}.docx")

# The ONE filter in this version. Rolling window from runtime, same
# convention as 01_ingest.py's DAYS_LOOKBACK.
DAYS_LOOKBACK = 15

SPOTLIGHT_COUNT = 3  # top N by disclosed value get full card treatment; the rest run as a compact list

FUNNEL_STAGES = [
    ("Raw ingested (RSS, last 30 days)", os.path.join("data", "raw_news.csv")),
    ("After deduplication", os.path.join("data", "processed_news.csv")),
    ("After relevance filter", os.path.join("data", "relevant_news.csv")),
    ("After credibility filter", os.path.join("data", "credible_news.csv")),
    (f"In this issue (last {DAYS_LOOKBACK} days)", None),  # filled in from final_df at runtime
]

PLURAL_LABELS = {
    "Acquisition": "Acquisition", "Investment": "Investment", "PE": "PE",
    "Funding": "Funding", "Merger": "Merger", "Other": "Other",
}

UNKNOWN_PLACEHOLDER = "Not specified"  # same convention as 05_summarize.py

# Approximate, documented assumption -- used only to make INR and USD
# deal values comparable for ranking/summing. Not a live exchange rate;
# update periodically if this matters for production use.
FX_RATE_INR_PER_USD = 87
SCALE_TO_MULTIPLIER = {
    "crore": 1e7, "cr": 1e7,
    "lakh": 1e5, "lakhs": 1e5,
    "million": 1e6, "mn": 1e6, "m": 1e6,
    "billion": 1e9, "bn": 1e9, "b": 1e9,
}
VALUE_PARSE_RE = re.compile(
    r"(₹|Rs\.?|INR|\$|USD)\s?([\d,]*\d(?:\.\d+)?)[\s-]?([A-Za-z]+)?",
    re.IGNORECASE,
)

# Single-accent palette. Charcoal does almost all the work; gold is the
# only accent and is used consistently for rules, emphasis, and values.
INK = RGBColor(0x1C, 0x1C, 0x1C)
BODY_GREY = RGBColor(0x44, 0x44, 0x44)
MUTED_GREY = RGBColor(0x88, 0x88, 0x88)
GOLD = RGBColor(0x9C, 0x7A, 0x32)
BOX_BORDER_GREY = "E5E5E5"
CREDIBILITY_COLORS = {
    "HIGH": RGBColor(0x2F, 0x6B, 0x3F),
    "MEDIUM": RGBColor(0xB7, 0x79, 0x0E),
    "LOW": RGBColor(0x6B, 0x6B, 0x6B),
}

SERIF_FONT = "Georgia"
SANS_FONT = "Arial"


# ─────────────────────────────────────────────────────────────────────────
# TEXT / DATE / VALUE HELPERS
# ─────────────────────────────────────────────────────────────────────────

def split_title(title):
    """Split an RSS title into (headline, publisher_suffix) on the LAST
    separator. Self-contained, same logic as 04/05."""
    title = str(title)
    parts = re.split(r"\s[-–—]\s", title)
    if len(parts) > 1:
        return " - ".join(parts[:-1]).strip(), parts[-1].strip()
    return title.strip(), ""


def clean_title(title):
    headline, _ = split_title(title)
    return headline


def parse_published_date(published_str):
    """Per-row RFC-822 parsing -- deliberately not pandas.to_datetime() on
    the whole column. Confirmed during 02_deduplicate.py testing that
    pandas' vectorized format inference silently drops mixed GMT/+0530
    timestamps, which this dataset has both of."""
    try:
        dt = parsedate_to_datetime(published_str)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc)
    except (TypeError, ValueError):
        return None


def parse_deal_value_to_usd_millions(value_text):
    """
    Normalize a deal_value string (e.g. '₹200 cr', '$4.3bn') into a
    comparable float, in USD millions, for ranking and summing. Returns
    None if the text is the unknown placeholder, or if it parses to a
    number but has no recognizable scale word -- a truncated upstream
    value like bare '$5' is discarded rather than guessed.
    """
    if not value_text or value_text == UNKNOWN_PLACEHOLDER:
        return None
    match = VALUE_PARSE_RE.match(str(value_text).strip())
    if not match:
        return None
    currency, number_str, scale_word = match.groups()
    try:
        number = float(number_str.replace(",", ""))
    except ValueError:
        return None
    multiplier = SCALE_TO_MULTIPLIER.get((scale_word or "").lower())
    if multiplier is None:
        return None
    is_inr = currency == "₹" or currency.lower().rstrip(".") in ("rs", "inr")
    value = number * multiplier
    value_usd = value / FX_RATE_INR_PER_USD if is_inr else value
    return value_usd / 1e6


def format_usd_millions(value_millions):
    if value_millions is None:
        return None
    if value_millions >= 1000:
        return f"${value_millions / 1000:.1f}B"
    return f"${value_millions:.0f}M"


# ─────────────────────────────────────────────────────────────────────────
# DATA SELECTION -- recency window only
# ─────────────────────────────────────────────────────────────────────────

def build_newsletter_pool(df):
    """
    The only filter applied: published within the last DAYS_LOOKBACK
    days (rolling from runtime). Every article that passes is included,
    ranked by disclosed value (largest first, unknown last) with recency
    as a tiebreaker -- this ordering decides Spotlight vs. ticker
    placement, it does not exclude anything.
    """
    pool = df.copy()
    pool["_published_dt"] = pool["published"].apply(parse_published_date)

    cutoff = datetime.now(timezone.utc) - timedelta(days=DAYS_LOOKBACK)
    pool = pool[pool["_published_dt"].apply(lambda d: d is not None and d >= cutoff)]

    if pool.empty:
        return pool

    pool["_usd_value"] = pool["deal_value"].apply(parse_deal_value_to_usd_millions)
    pool["_has_value"] = pool["_usd_value"].notna()

    pool = pool.sort_values(
        ["_has_value", "_usd_value", "_published_dt"],
        ascending=[False, False, False],
        na_position="last",
    )
    return pool.reset_index(drop=True)


# ─────────────────────────────────────────────────────────────────────────
# EXECUTIVE HIGHLIGHTS (data-derived, no LLM, no invented commentary)
# ─────────────────────────────────────────────────────────────────────────

def build_executive_highlights(final_df):
    bullets = []
    if final_df.empty:
        return bullets

    valued = final_df[final_df["_usd_value"].notna()]
    if not valued.empty:
        top = valued.iloc[0]
        acquirer, target = top.get("acquirer", ""), top.get("target", "")
        if acquirer != UNKNOWN_PLACEHOLDER and target != UNKNOWN_PLACEHOLDER:
            subject = f"{acquirer}'s {top['deal_value']} acquisition of {target}"
        else:
            subject = f"{clean_title(top['title'])} ({top['deal_value']})"
        bullets.append(f"{subject} was the largest disclosed deal this issue.")

    total_value = valued["_usd_value"].sum() if not valued.empty else None
    if total_value:
        bullets.append(
            f"{len(final_df)} articles tracked this issue, with a combined disclosed "
            f"value of {format_usd_millions(total_value)} across {len(valued)} item(s)."
        )
    else:
        bullets.append(f"{len(final_df)} articles tracked this issue.")

    type_counts = final_df["deal_type"].value_counts()
    if not type_counts.empty:
        top_type, top_count = type_counts.index[0], int(type_counts.iloc[0])
        bullets.append(
            f"{PLURAL_LABELS.get(top_type, top_type)} was the most common deal type "
            f"this cycle, with {top_count} of {len(final_df)} tracked items."
        )

    known_acquirers = final_df[final_df["acquirer"] != UNKNOWN_PLACEHOLDER]["acquirer"]
    acquirer_counts = known_acquirers.value_counts()
    if not acquirer_counts.empty and acquirer_counts.iloc[0] >= 2:
        bullets.append(
            f"{acquirer_counts.index[0]} was the most active acquirer this issue, "
            f"appearing in {int(acquirer_counts.iloc[0])} tracked items."
        )

    return bullets[:4]


def get_funnel_counts(final_df):
    counts = []
    for label, path in FUNNEL_STAGES:
        if path is None:
            counts.append((label, len(final_df)))
            continue
        n = None
        if os.path.exists(path):
            try:
                n = len(pd.read_csv(path))
            except Exception:
                n = None
        counts.append((label, n))
    return counts


# ─────────────────────────────────────────────────────────────────────────
# DOCX LOW-LEVEL HELPERS
# ─────────────────────────────────────────────────────────────────────────

def fix_zoom_settings(doc):
    """python-docx's default template ships a <w:zoom> element missing a
    schema-required attribute -- confirmed by running the OOXML validator
    against a blank python-docx document. Word opens the file fine
    regardless, but this gets a clean strict-schema validation pass."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def add_hyperlink(paragraph, url, text, color="9C7A32", underline=True, size=9):
    """Insert a clickable hyperlink run. python-docx 1.2.0 only exposes a
    read-only `paragraph.hyperlinks` property, so this builds the
    required <w:hyperlink> XML directly."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_element = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    font_el = OxmlElement("w:rFonts")
    font_el.set(qn("w:ascii"), SANS_FONT)
    run_props.append(font_el)

    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(size * 2))
    run_props.append(size_el)

    if color:
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color)
        run_props.append(color_el)
    if underline:
        underline_el = OxmlElement("w:u")
        underline_el.set(qn("w:val"), "single")
        run_props.append(underline_el)

    run_element.append(run_props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_element.append(text_el)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


# CT_PPr requires its children in this exact schema order. Blindly
# appending a new child (e.g. w:pBdr) after python-docx has already
# inserted w:spacing produces an invalid document -- confirmed by running
# the strict OOXML validator against this script's own output.
PPR_CHILD_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
    "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind", "contextualSpacing",
    "mirrorIndents", "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange",
]


def insert_pPr_child(pPr, new_element, tag):
    target_idx = PPR_CHILD_ORDER.index(tag)
    for child in pPr:
        child_tag = child.tag.split("}")[-1]
        if child_tag in PPR_CHILD_ORDER and PPR_CHILD_ORDER.index(child_tag) > target_idx:
            child.addprevious(new_element)
            return
    pPr.append(new_element)


# CT_TblPrBase requires its children in this exact schema order too --
# table.autofit = False causes python-docx to insert w:tblLayout into
# tblPr; appending w:tblBorders after that (it belongs BEFORE tblLayout)
# produced the same class of validation error as the pPr case above.
TBLPR_CHILD_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
    "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
]


def insert_tblPr_child(tblPr, new_element, tag):
    target_idx = TBLPR_CHILD_ORDER.index(tag)
    for child in tblPr:
        child_tag = child.tag.split("}")[-1]
        if child_tag in TBLPR_CHILD_ORDER and TBLPR_CHILD_ORDER.index(child_tag) > target_idx:
            child.addprevious(new_element)
            return
    tblPr.append(new_element)


def add_rule(doc, color=GOLD, weight_pt=1.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(weight_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    border.append(bottom)
    insert_pPr_child(pPr, border, "pBdr")
    return p


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def set_table_borders(table, color_hex=BOX_BORDER_GREY, size=4, left_accent=None):
    """Apply a real box border to a table (the 'outline box' for each
    news item). size is in eighths of a point. If left_accent is given,
    the left border is thicker and gold."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        if edge == "left" and left_accent:
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), left_accent)
        else:
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color_hex)
        el.set(qn("w:space"), "0")
        borders.append(el)
    insert_tblPr_child(tblPr, borders, "tblBorders")


def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Internal padding for a table cell, in twips. Element order matters
    -- CT_TcMar's schema sequence is top, start, left, bottom, end, right
    (confirmed by reading wml.xsd directly), so top/left/bottom/right in
    that order is required; left must come before bottom."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def make_box(doc, left_accent=None):
    """Create a single-cell table that acts as a bordered, padded box --
    the 'outline box' for one news item. Returns the cell to add content
    paragraphs into."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_table_borders(table, left_accent=left_accent)
    set_cell_margins(cell)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    return cell


# ─────────────────────────────────────────────────────────────────────────
# DOCUMENT BUILDING
# ─────────────────────────────────────────────────────────────────────────

def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = SANS_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_GREY


def add_masthead(doc, issue_label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("FMCG M&A Intelligence")
    run.font.name = SERIF_FONT
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = INK

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"{issue_label}   •   Prepared for Executive Leadership")
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = MUTED_GREY

    add_rule(doc)


def add_contents_strip(doc, final_df):
    p = doc.add_paragraph()
    label_run = p.add_run("THIS ISSUE   ")
    label_run.font.size = Pt(8.5)
    label_run.font.color.rgb = MUTED_GREY
    content_run = p.add_run(
        f"{len(final_df)} articles  •  last {DAYS_LOOKBACK} days  •  ranked by disclosed value"
    )
    content_run.font.size = Pt(10)
    content_run.font.bold = True
    content_run.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)
    add_rule(doc, weight_pt=0.75, color=RGBColor(0xE5, 0xE5, 0xE5))


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    add_rule(doc, weight_pt=2, space_after=8)


def add_executive_highlights(doc, bullets):
    add_section_header(doc, "EXECUTIVE HIGHLIGHTS")
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(bullet)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)



def add_deal_box_content(cell, row, compact=False):
    """Populate a boxed cell (spotlight or ticker) with one item's
    content. `compact=True` produces the single-paragraph ticker style;
    `compact=False` produces the full multi-paragraph spotlight card."""
    credibility = str(row.get("credibility", "MEDIUM")).upper()
    cred_color = CREDIBILITY_COLORS.get(credibility, MUTED_GREY)
    acquirer, target = str(row.get("acquirer", "")), str(row.get("target", ""))
    deal_value = str(row.get("deal_value", UNKNOWN_PLACEHOLDER))
    deal_type_label = PLURAL_LABELS.get(str(row.get("deal_type", "Other")), "Other")
    headline = (f"{acquirer} \u2192 {target}"
                if acquirer != UNKNOWN_PLACEHOLDER and target != UNKNOWN_PLACEHOLDER
                else clean_title(row["title"]))
    publisher = str(row.get("publisher", "unknown")).title()
    pub_dt = row.get("_published_dt")

    if compact:
        p = cell.paragraphs[0]
        date_str = pub_dt.strftime("%d %b") if pub_dt is not None else ""
        synopsis = str(row.get("synopsis", "")).strip() or clean_title(row["title"])
        value_text = f"  \u00B7  {deal_value}" if deal_value != UNKNOWN_PLACEHOLDER else ""

        headline_run = p.add_run(headline)
        headline_run.font.bold = True
        headline_run.font.size = Pt(10.5)
        headline_run.font.color.rgb = INK

        detail_run = p.add_run(f"{value_text}  \u00B7  {synopsis}  \u00B7  {deal_type_label}  \u00B7  ")
        detail_run.font.size = Pt(9.5)
        detail_run.font.color.rgb = MUTED_GREY

        dot_run = p.add_run("\u25CF ")
        dot_run.font.size = Pt(7)
        dot_run.font.color.rgb = cred_color
        cred_run = p.add_run(f"{credibility}  \u00B7  {publisher}, {date_str}")
        cred_run.font.size = Pt(9.5)
        cred_run.font.color.rgb = MUTED_GREY
        return

    headline_p = cell.paragraphs[0]
    headline_run = headline_p.add_run(headline)
    headline_run.font.bold = True
    headline_run.font.size = Pt(13.5)
    headline_run.font.color.rgb = INK
    if deal_value != UNKNOWN_PLACEHOLDER:
        value_run = headline_p.add_run(f"    {deal_value}")
        value_run.font.bold = True
        value_run.font.size = Pt(12)
        value_run.font.color.rgb = GOLD

    tag_p = cell.add_paragraph()
    tag_p.paragraph_format.space_before = Pt(1)
    tag_p.paragraph_format.space_after = Pt(3)
    tag_run = tag_p.add_run(deal_type_label.upper())
    tag_run.font.size = Pt(8)
    tag_run.font.bold = True
    tag_run.font.color.rgb = GOLD

    synopsis = str(row.get("synopsis", "")).strip() or clean_title(row["title"])
    synopsis_p = cell.add_paragraph()
    synopsis_p.paragraph_format.space_after = Pt(6)
    synopsis_run = synopsis_p.add_run(synopsis)
    synopsis_run.font.size = Pt(11)
    synopsis_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    source_p = cell.add_paragraph()
    source_p.paragraph_format.space_after = Pt(0)
    source_run = source_p.add_run(f"{publisher}   ")
    source_run.font.size = Pt(9)
    source_run.font.color.rgb = MUTED_GREY

    dot_run = source_p.add_run("\u25CF ")
    dot_run.font.size = Pt(8)
    dot_run.font.color.rgb = cred_color
    cred_run = source_p.add_run(f"{credibility}   ")
    cred_run.font.bold = True
    cred_run.font.size = Pt(9)
    cred_run.font.color.rgb = cred_color

    date_str = pub_dt.strftime("%d %b %Y") if pub_dt is not None else "date unknown"
    date_run = source_p.add_run(f"{date_str}   ")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = MUTED_GREY

    if str(row.get("url", "")).startswith("http"):
        add_hyperlink(source_p, str(row["url"]), "Read full article")


def add_spotlight_card(doc, row):
    cell = make_box(doc, left_accent="9C7A32")
    add_deal_box_content(cell, row, compact=False)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_ticker_box(doc, row):
    cell = make_box(doc)
    add_deal_box_content(cell, row, compact=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_appendix(doc, final_df):
    doc.add_page_break()
    add_section_header(doc, "APPENDIX — PIPELINE METHODOLOGY")
    doc.add_paragraph(
        "Every article in this issue passed through the standard automated pipeline: "
        "RSS ingestion (30-day recency window), near-duplicate removal (exact match + "
        "semantic similarity), FMCG-relevance filtering, and publisher-credibility "
        f"scoring. This newsletter additionally restricts to articles published in the "
        f"last {DAYS_LOOKBACK} days, ranked by disclosed value (largest first); no "
        "further filtering is applied at this stage."
    )

    counts = get_funnel_counts(final_df)
    raw_count = counts[0][1]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, heading in enumerate(["Pipeline stage", "Articles", "% of raw ingested"]):
        header_cells[i].text = heading
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_background(header_cells[i], "D9D9D9")

    for label, count in counts:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(count) if count is not None else "N/A"
        if count is not None and raw_count:
            row_cells[2].text = f"{count / raw_count * 100:.0f}%"
        else:
            row_cells[2].text = "N/A"


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""

    left_run = footer_p.add_run(
        "Automated screening pipeline: ingestion \u2192 dedup \u2192 relevance \u2192 "
        "credibility \u2192 summarization.\t"
    )
    left_run.font.size = Pt(7.5)
    left_run.font.color.rgb = MUTED_GREY

    add_page_number_field(footer_p)
    for run in footer_p.runs[-4:]:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED_GREY

    tab_stops = footer_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))


# ─────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────

def main():
    print("\n" + "=" * 70)
    print(f"STAGE 6 - NEWSLETTER GENERATION   |   last {DAYS_LOOKBACK} days")
    print("=" * 70 + "\n")

    if not os.path.exists(INPUT_FILE):
        print(f"Input file not found: {INPUT_FILE}")
        print("Run 05_summarize.py first.\n")
        return

    df = pd.read_csv(INPUT_FILE, encoding="utf-8")
    if df.empty:
        print("Input file has 0 rows -- nothing to put in the newsletter.\n")
        return

    df["title"] = df["title"].fillna("")
    df["summary"] = df["summary"].fillna("")
    df["deal_type"] = df["deal_type"].fillna("Other")
    df["acquirer"] = df["acquirer"].fillna(UNKNOWN_PLACEHOLDER)
    df["target"] = df["target"].fillna(UNKNOWN_PLACEHOLDER)

    pre_filter_count = len(df)
    final_df = build_newsletter_pool(df)

    if final_df.empty:
        print(f"No articles published in the last {DAYS_LOOKBACK} days. Nothing generated.\n")
        return

    doc = Document()
    fix_zoom_settings(doc)
    set_default_style(doc)

    issue_label = f"Issue: {datetime.now().strftime('%B %Y')}"
    add_masthead(doc, issue_label)
    add_contents_strip(doc, final_df)

    bullets = build_executive_highlights(final_df)
    add_executive_highlights(doc, bullets)

    spotlight = final_df.head(SPOTLIGHT_COUNT)
    ticker = final_df.iloc[SPOTLIGHT_COUNT:]

    add_section_header(doc, f"SPOTLIGHT — TOP {len(spotlight)} BY DISCLOSED VALUE")
    for _, row in spotlight.iterrows():
        add_spotlight_card(doc, row)

    if not ticker.empty:
        add_section_header(doc, "ALSO THIS ISSUE")
        for _, row in ticker.iterrows():
            add_ticker_box(doc, row)

    add_appendix(doc, final_df)
    add_footer(doc)

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)

    # ── Run summary printed to console ─────────────────────────────────
    valued = final_df[final_df["_usd_value"].notna()]
    total_value = valued["_usd_value"].sum() if not valued.empty else None

    print("-" * 70)
    print("NEWSLETTER SUMMARY")
    print("-" * 70)
    print(f"Rows in summarized_news.csv                 : {pre_filter_count}")
    print(f"Articles in last {DAYS_LOOKBACK} days                      : {len(final_df)}")
    print(f"  Credibility breakdown                     : "
          f"{dict(final_df['credibility'].value_counts())}")
    print(f"Articles with a disclosed value              : {len(valued)}")
    print(f"Total disclosed value                       : {format_usd_millions(total_value) or 'N/A'}")
    print(f"Spotlight (full card)                       : {len(spotlight)}")
    print(f"Also this issue (compact list)               : {len(ticker)}")
    print(f"Output file                                 : {OUTPUT_FILE}")
    print("-" * 70 + "\n")
if __name__ == "__main__":
    main()